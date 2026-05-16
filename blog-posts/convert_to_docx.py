"""Convert blog markdown to docx for review."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

with open('BLOG-SRE-AGENT-PRIVATE-AKS-DRAFT.md', 'r') as f:
    lines = f.readlines()


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text.rstrip())
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2e)


def add_table_from_lines(doc, table_lines):
    rows = []
    for tl in table_lines:
        tl = tl.strip()
        if tl.startswith('|'):
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            if cells and all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style='Light Grid Accent 1')
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell_text = cell_text.replace('**', '')
                table.rows[ri].cells[ci].text = cell_text
                if ri == 0:
                    for paragraph in table.rows[ri].cells[ci].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def add_paragraph_with_formatting(doc, text):
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean = part[2:-2].replace('`', '')
            run = p.add_run(clean)
            run.bold = True
        else:
            code_parts = re.split(r'(`.+?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = p.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    cp = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', cp)
                    run = p.add_run(cp)
    return p


i = 0
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith('```') and not in_code_block:
        in_code_block = True
        code_lines = []
        i += 1
        continue
    elif line.strip().startswith('```') and in_code_block:
        in_code_block = False
        add_code_block(doc, '\n'.join(code_lines))
        code_lines = []
        i += 1
        continue
    elif in_code_block:
        code_lines.append(line.rstrip())
        i += 1
        continue

    stripped = line.strip()

    # Skip HTML comments
    if stripped.startswith('<!--') and '-->' in stripped:
        i += 1
        continue

    # Skip empty lines
    if not stripped:
        i += 1
        continue

    # Horizontal rules
    if stripped == '---':
        i += 1
        continue

    # Headers
    if stripped.startswith('# ') and not stripped.startswith('## '):
        doc.add_heading(stripped[2:], level=1)
        i += 1
        continue
    elif stripped.startswith('## ') and not stripped.startswith('### '):
        doc.add_heading(stripped[3:], level=2)
        i += 1
        continue
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
        i += 1
        continue

    # Tables
    if stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        continue

    # Blockquotes
    if stripped.startswith('> '):
        text = stripped[2:]
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = text.replace('`', '')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # Bullet points
    if stripped.startswith('- ['):
        # Checklist items
        text = re.sub(r'- \[.\] ', '', stripped)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = text.replace('`', '')
        doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue
    elif stripped.startswith('- '):
        text = stripped[2:]
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = text.replace('`', '')
        doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue

    # Figure captions
    if stripped.startswith('*Figure') or stripped.startswith('*Tags:'):
        text = stripped.strip('*')
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Italic lines
    if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
        text = stripped.strip('*')
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        i += 1
        continue

    # Regular paragraphs
    add_paragraph_with_formatting(doc, stripped)
    i += 1

doc.save('BLOG-SRE-AGENT-PRIVATE-AKS-DRAFT.docx')
print('Done: BLOG-SRE-AGENT-PRIVATE-AKS-DRAFT.docx')
